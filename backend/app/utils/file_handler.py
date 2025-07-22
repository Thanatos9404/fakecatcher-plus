from fastapi import UploadFile
import fitz  # PyMuPDF
import docx
import io
import logging
from typing import Optional
from PIL import Image
import pytesseract

logger = logging.getLogger(__name__)


class FileHandler:
    """Handle file operations and text extraction for resumes and job postings"""

    def __init__(self):
        # Configure OCR settings for job posting images
        self.ocr_config = r'--oem 3 --psm 6 -l eng'
        logger.info("FileHandler initialized with OCR support")

    async def extract_text_from_resume(self, file: UploadFile) -> str:
        """Extract text content from resume file (PDF/DOC/DOCX)"""

        if not file.filename:
            raise ValueError("No filename provided")

        file_extension = file.filename.lower().split('.')[-1]
        logger.info(f"Processing resume file with extension: {file_extension}")

        try:
            if file_extension == 'pdf':
                return await self._extract_from_pdf_pymupdf(file)
            elif file_extension in ['doc', 'docx']:
                return await self._extract_from_docx(file)
            else:
                raise ValueError(f"Unsupported resume file format: {file_extension}")

        except Exception as e:
            logger.error(f"Failed to extract text from resume {file_extension}: {str(e)}")
            raise Exception(f"Failed to extract text from resume: {str(e)}")

    async def extract_text_from_job_posting(self, file: UploadFile, file_type: str) -> str:
        """Extract text from job posting (images or PDFs)"""

        if not file.filename:
            raise ValueError("No filename provided")

        logger.info(f"Processing job posting file: {file.filename}, type: {file_type}")

        try:
            if file_type == 'image':
                return await self._extract_from_image_ocr(file)
            elif file_type == 'pdf':
                return await self._extract_from_pdf_pymupdf(file)
            else:
                raise ValueError(f"Unsupported job posting file type: {file_type}")

        except Exception as e:
            logger.error(f"Failed to extract text from job posting: {str(e)}")
            raise Exception(f"Failed to extract text from job posting: {str(e)}")

    async def _extract_from_image_ocr(self, file: UploadFile) -> str:
        """Extract text from image using OCR"""
        try:
            # Read image content
            content = await file.read()
            logger.info(f"Read image content, size: {len(content)} bytes")

            if len(content) == 0:
                raise ValueError("Image file is empty")

            # Open image with PIL
            image = Image.open(io.BytesIO(content))
            logger.info(f"Image opened successfully, size: {image.size}, mode: {image.mode}")

            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')

            # Perform OCR
            extracted_text = pytesseract.image_to_string(image, config=self.ocr_config)

            # Reset file position
            await file.seek(0)

            logger.info(f"OCR extraction completed, text length: {len(extracted_text)}")

            if not extracted_text.strip():
                raise ValueError("No text could be extracted from the image")

            return extracted_text.strip()

        except Exception as e:
            logger.error(f"OCR extraction failed: {str(e)}")
            await file.seek(0)
            raise Exception(f"Could not extract text from image: {str(e)}")

    async def _extract_from_pdf_pymupdf(self, file: UploadFile) -> str:
        """Extract text from PDF using PyMuPDF"""
        try:
            # Read file content
            content = await file.read()
            logger.info(f"Read PDF content, size: {len(content)} bytes")

            if len(content) == 0:
                raise ValueError("PDF file is empty")

            # Open PDF from memory
            pdf_doc = fitz.open(stream=content, filetype="pdf")
            logger.info(f"PDF opened successfully, pages: {pdf_doc.page_count}")

            text = ""
            for page_num in range(pdf_doc.page_count):
                try:
                    page = pdf_doc[page_num]
                    page_text = page.get_text()
                    if page_text:
                        text += page_text + "\n"
                        logger.info(f"Extracted text from page {page_num + 1}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")

            pdf_doc.close()

            # Reset file position
            await file.seek(0)

            logger.info(f"Total extracted text length: {len(text)}")

            if not text.strip():
                raise ValueError("No text could be extracted from the PDF")

            return text.strip()

        except Exception as e:
            logger.error(f"PyMuPDF extraction failed: {str(e)}")
            await file.seek(0)
            raise Exception(f"Could not extract text from PDF: {str(e)}")

    async def _extract_from_docx(self, file: UploadFile) -> str:
        """Extract text from DOCX file"""
        try:
            content = await file.read()
            logger.info(f"Read DOCX content, size: {len(content)} bytes")

            if len(content) == 0:
                raise ValueError("DOCX file is empty")

            # Parse DOCX
            doc = docx.Document(io.BytesIO(content))

            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            extracted_text = '\n'.join(text_parts)

            # Reset file position
            await file.seek(0)

            logger.info(f"DOCX extraction completed, text length: {len(extracted_text)}")

            if not extracted_text.strip():
                raise ValueError("No text could be extracted from the DOCX file")

            return extracted_text.strip()

        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            await file.seek(0)
            raise Exception(f"Could not extract text from DOCX: {str(e)}")
